from openai import OpenAI
from docx import Document
import os
import streamlit as st

def enhance_document_language(docx_path, output_path):
    # Load the document
    doc = Document(docx_path)
    enhanced_doc = Document()
    
    client = OpenAI(
    api_key = st.secrets["OPENAI_API_KEY"]
    )

    # Function to enhance text
    def enhance_text(text):
        try:
            response = client.completions.create(
                model = "gpt-3.5-turbo-instruct",
                prompt=text,
                max_tokens=1500,
                n=1,
                stop=None,
                temperature=0.7
            )
            return response.choices[0].text.strip()
        except Exception as e:
            print(f"Failed to enhance text: {str(e)}")
            return text  # Return the original text if the API call fails

    # Read and enhance each paragraph
    for para in doc.paragraphs:
        enhanced_text = enhance_text(para.text)
        enhanced_doc.add_paragraph(enhanced_text)

    # Save the enhanced document
    enhanced_doc.save(output_path)

enhance_document_language('input.docx', 'enhanced_output.docx')
