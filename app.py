import streamlit as st
from openai import OpenAI
from docx import Document
from io import BytesIO
from docx.shared import Pt


def modify_docx_styles(doc, output_path):
    
    for para in doc.paragraphs:
        
        # Check if the paragraph style is a heading style
        if para.style.name.startswith('Heading'):
            # Set font size for headings
            for run in para.runs:
                run.font.size = Pt(36)  # 36 points in half-points (36 * 20)
        else:
             for run in para.runs:
                run.font.size = Pt(14)
                
        # Set the font style for all text
        for run in para.runs:
            run.font.name = "EYInterstate Light"
            # run.font.size = Pt(14)
    
    # Save the modified document
    doc.save(output_path)


def load_docx(file):
    """Load DOCX file into a Document object."""
    return Document(file)

def docx_to_text(doc):
    """Extract text from DOCX Document."""
    return '\n\n'.join([para.text for para in doc.paragraphs])

def text_to_docx(text):
    """Convert text back into a DOCX Document."""
    doc = Document()
    for line in text.split('\n'):
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(line)
        run.font.name = 'EYInterstate Light'
        
    return doc

def improve_text_with_openai(text, api_key):
    
    client = OpenAI(
    api_key = api_key,
    )
    response = client.chat.completions.create(
                    model = "gpt-3.5-turbo",
                      messages=[
    {"role": "system", "content": "You are a helpful assistant. Rewrite the following text in a business language"},
    {"role": "user", "content": text}
  ]
                    )
    return response.choices[0].message.content

def save_docx(doc):
    """Save Document object to a BytesIO object and return it."""
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def extract_font_styles(document):
   
    # Set to collect unique font styles
    font_styles = set()
    
    # Iterate through each paragraph and run in the document
    for para in document.paragraphs:
        for run in para.runs:
            if run.font.name is not None:
                font_styles.add(run.font.name)

    # Also check all tables
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.font.name is not None:
                            font_styles.add(run.font.name)
    
    return font_styles
    

def main():
    st.title("DOCX Enhancer")

    uploaded_file = st.file_uploader("Upload your file", type=['docx'])
    api_key = st.secrets["OPENAI_API_KEY"]

    if uploaded_file and api_key:
        with st.spinner('Working on your document...'):
            doc = load_docx(uploaded_file)
            font_styles = extract_font_styles(doc)
            if(not(len(font_styles) == 1 and "EYInterstate Light" in font_styles)):
                st.error("Your document uses the following Font styles other than EY Interstate - " + ', '.join(font_styles) + ". Converting the document to EY Interstate Font and fixing the font size.")
                   
            original_text = docx_to_text(doc)
            improved_text = improve_text_with_openai(original_text, api_key)
            # st.write(improved_text)
            new_doc = text_to_docx(improved_text)
            buffer = save_docx(new_doc)

            st.success("Enhancement Complete! Download your improved DOCX below.")
            st.download_button(label="Download DOCX", data=buffer, file_name="enhanced_file.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

if __name__ == "__main__":
    main()
